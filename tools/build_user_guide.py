from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
OUT = DOCS / "DersDagitim_Kullanim_Kilavuzu.docx"

SCREENSHOTS = [
    (
        Path(r"C:\Users\masco\AppData\Local\Temp\codex-clipboard-ba6c68f3-1314-44a3-bfb2-f35c12554e20.png"),
        "Taslak çizelge: okul geneli ders yoğunluğu, seçili atama detayı ve çıktı butonları.",
    ),
    (
        Path(r"C:\Users\masco\AppData\Local\Temp\codex-clipboard-dd294424-073d-43d5-9a68-2d76f2481a44.png"),
        "Çakışma kontrolü: yerleşmeyen veya gerçek bindirme varsa burada listelenir.",
    ),
    (
        Path(r"C:\Users\masco\AppData\Local\Temp\codex-clipboard-3aa81e8e-cf86-4ca5-aa80-54b4a7610748.png"),
        "Öğretmen programı ve ders yükü: öğretmen bazlı ders dağılımı ve görev bilgisi.",
    ),
    (
        Path(r"C:\Users\masco\AppData\Local\Temp\codex-clipboard-a700e071-8c0f-45a3-bbfc-1eb619832205.png"),
        "Öğretmen uygunluk kilitleri: öğretmenlerin kesin kilit veya uygunluk kayıtları.",
    ),
]


def font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, box: tuple[int, int, int, int], fnt, fill, line_gap=6):
    x1, y1, x2, _ = box
    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        trial = f"{line} {word}".strip()
        width = draw.textbbox((0, 0), trial, font=fnt)[2]
        if width <= x2 - x1 or not line:
            line = trial
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    y = y1
    for line in lines:
        draw.text((x1, y), line, font=fnt, fill=fill)
        y += fnt.size + line_gap


def create_flow_image(path: Path):
    img = Image.new("RGB", (1400, 760), "#F8FAFC")
    d = ImageDraw.Draw(img)
    title = font(42, True)
    body = font(24)
    small = font(19)
    d.text((52, 38), "Ders Dağıtım Uygulaması - temel iş akışı", font=title, fill="#0F172A")

    steps = [
        ("1", "Tanımlar", "ASC XML içe aktar; sınıf, öğretmen, ders ve kaynak kayıtlarını kontrol et."),
        ("2", "Kurallar", "Öğretmen uygunluk ve kesin kilit kayıtlarını gözden geçir."),
        ("3", "Taslak Üret", "Mevcut ASC kartlarını koruyarak 522 talep / 1259 kart üzerinden taslak oluştur."),
        ("4", "Manuel Düzenle", "Dersi gün/saat olarak taşı, gerekirse atamayı kaldır ve geri al."),
        ("5", "Kontrol ve Çıktı", "Çakışmaları incele; sınıf, kaynak, öğretmen ve okul geneli çıktı al."),
    ]

    x = 52
    y = 130
    card_w = 240
    gap = 22
    for i, (num, heading, desc) in enumerate(steps):
        d.rounded_rectangle((x, y, x + card_w, y + 430), radius=18, fill="#FFFFFF", outline="#CBD5E1", width=2)
        d.ellipse((x + 22, y + 22, x + 82, y + 82), fill="#2E74B5")
        d.text((x + 45, y + 35), num, font=font(25, True), fill="#FFFFFF", anchor="mm")
        d.text((x + 22, y + 110), heading, font=font(25, True), fill="#0F172A")
        draw_wrapped(d, desc, (x + 22, y + 160, x + card_w - 22, y + 360), small, "#334155")
        if i < len(steps) - 1:
            d.line((x + card_w + 2, y + 210, x + card_w + gap - 4, y + 210), fill="#64748B", width=4)
            d.polygon(
                [
                    (x + card_w + gap - 4, y + 210),
                    (x + card_w + gap - 18, y + 200),
                    (x + card_w + gap - 18, y + 220),
                ],
                fill="#64748B",
            )
        x += card_w + gap

    d.rounded_rectangle((52, 600, 1348, 705), radius=16, fill="#EFF6FF", outline="#BFDBFE", width=2)
    draw_wrapped(
        d,
        "Önemli ilke: Program otomatik taslak üretir; manuel gün/saat değişiklikleri SQLite içinde korunur ve sonraki taslak üretiminde dikkate alınır.",
        (82, 626, 1320, 690),
        body,
        "#1E3A8A",
    )
    img.save(path)


def create_manual_panel_image(path: Path):
    img = Image.new("RGB", (1200, 720), "#F8FAFC")
    d = ImageDraw.Draw(img)
    d.text((50, 38), "Taslak ekranında manuel düzenleme nerede?", font=font(38, True), fill="#0F172A")

    d.rounded_rectangle((50, 110, 815, 620), radius=12, fill="#FFFFFF", outline="#CBD5E1", width=2)
    d.text((80, 135), "Haftalık çizelge", font=font(26, True), fill="#334155")
    days = ["Pzt", "Sal", "Çar", "Per", "Cum"]
    for i, day in enumerate(days):
        x = 165 + i * 125
        d.rectangle((x, 185, x + 120, 225), fill="#E8EEF5", outline="#CBD5E1")
        d.text((x + 60, 207), day, font=font(18, True), fill="#334155", anchor="mm")
    for hour in range(1, 8):
        y = 225 + (hour - 1) * 48
        d.rectangle((80, y, 160, y + 46), fill="#F1F5F9", outline="#CBD5E1")
        d.text((120, y + 23), str(hour), font=font(17, True), fill="#475569", anchor="mm")
        for i in range(5):
            x = 165 + i * 125
            fill = "#DBEAFE" if (hour, i) in [(2, 1), (4, 3), (6, 0)] else "#FFFFFF"
            if (hour, i) == (4, 3):
                fill = "#B7E4C7"
            d.rectangle((x, y, x + 120, y + 46), fill=fill, outline="#E2E8F0")
            if (hour, i) == (4, 3):
                d.text((x + 60, y + 23), "11/D\nMobil", font=font(14, True), fill="#0F172A", anchor="mm")

    d.rounded_rectangle((860, 110, 1150, 620), radius=12, fill="#EFF6FF", outline="#93C5FD", width=3)
    d.text((885, 135), "Manuel düzenleme", font=font(24, True), fill="#1E3A8A")
    draw_wrapped(
        d,
        "1. Tek ders içeren hücreyi seç. 2. Yeni gün ve ders saatini belirle. 3. Seçili dersi bu saate taşı.",
        (885, 178, 1125, 260),
        font(17),
        "#334155",
    )
    d.text((885, 305), "Yeni gün", font=font(18, True), fill="#334155")
    d.rounded_rectangle((885, 335, 1125, 380), radius=8, fill="#FFFFFF", outline="#CBD5E1")
    d.text((905, 358), "Perşembe", font=font(18), fill="#0F172A", anchor="lm")
    d.text((885, 405), "Yeni ders saati", font=font(18, True), fill="#334155")
    d.rounded_rectangle((885, 435, 1125, 480), radius=8, fill="#FFFFFF", outline="#CBD5E1")
    d.text((905, 458), "5", font=font(18), fill="#0F172A", anchor="lm")
    d.rounded_rectangle((885, 515, 1125, 565), radius=8, fill="#2E74B5")
    d.text((1005, 540), "Seçili dersi taşı", font=font(18, True), fill="#FFFFFF", anchor="mm")

    d.line((785, 420, 860, 420), fill="#2563EB", width=6)
    d.polygon([(860, 420), (840, 407), (840, 433)], fill="#2563EB")
    img.save(path)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_steps(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_note(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EFF6FF")
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(30, 58, 138)
    table.autofit = False
    table.columns[0].width = Inches(6.3)


def add_image(doc: Document, path: Path, caption: str, width_in: float = 6.25):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].italic = True
    c.runs[0].font.size = Pt(9)
    c.runs[0].font.color.rgb = RGBColor(71, 85, 105)


def create_document():
    DOCS.mkdir(exist_ok=True)
    ASSETS.mkdir(exist_ok=True)

    flow_img = ASSETS / "kullanim_akisi.png"
    manual_img = ASSETS / "manuel_duzenleme_paneli.png"
    create_flow_image(flow_img)
    create_manual_panel_image(manual_img)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Ders Dağıtım Uygulaması\nKullanım Kılavuzu")
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(15, 23, 42)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Sıfırdan kullanacak idareci için adım adım rehber").italic = True
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Hazırlayan: Mehmet Akif SÖNMEZ | İkizsoft Bilişim Hizmetleri | www.ikizsoft.com | Sürüm: 0.9.0-dev")

    add_note(
        doc,
        "Bu kılavuz mevcut uygulama durumuna göre hazırlanmıştır. Program gerçek ASC XML kurum verisini yerel SQLite içinde saklar; taslak üretirken mevcut ASC kartlarını ve manuel gün/saat değişikliklerini korur.",
    )
    add_image(doc, flow_img, "Şekil 1 - Uygulamanın temel iş akışı")

    add_heading(doc, "1. Uygulama Ne İşe Yarar?", 1)
    doc.add_paragraph(
        "Bu uygulama okulun ders programı verisini ASC XML üzerinden içe alır, yerel SQLite veritabanında saklar, açıklanabilir taslak üretir, çakışma kontrolü yapar ve öğretmen, sınıf, kaynak/laboratuvar, okul geneli rapor çıktıları üretir."
    )
    add_bullets(
        doc,
        [
            "Ana hedef: ASC verisini bozmadan programı görmek, kontrol etmek, manuel gün/saat değişiklikleri yapmak ve çıktı almak.",
            "Manuel değişiklikler: Mevcut ASC kartlarına bağlı gün/saat taşıma ve kaldırma işlemleri kalıcı saklanır.",
            "Çıktılar: CSV, PDF/yazdırma, renkli okul dosyası ve ASC XML dışa aktarım.",
        ],
    )

    add_heading(doc, "2. İlk Açılış ve Ana Menü", 1)
    doc.add_paragraph("Ana ekrandaki sol menü iş akışını sırayla verir. İlk kez kullanan biri aşağıdaki sırayı takip etmelidir.")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Menü", "Ne İçin Kullanılır?", "İlk kullanımda yapılacak işlem"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    rows = [
        ("1 - Tanımlar", "Sınıf, öğretmen, ders, kaynak ve ASC XML işlemleri", "ASC XML içe aktar; kayıtları kontrol et."),
        ("2 - Kurallar ve Kilitler", "Öğretmen uygunluk / kesin kilit bilgisi", "Kısıt varsa kontrol et."),
        ("3 - Taslak Üret", "Gerçek ASC taleplerinden program taslağı", "Mevcut ASC kartlarını koru seçeneğiyle üret."),
        ("4 - Çakışma Kontrolü", "Yerleşmeyen ve gerçek bindirme analizi", "Taslak sonrası kontrol et."),
        ("5 - Çizelgeler", "Sınıf ve kaynak/lab programları", "Sınıf veya kaynak seçip çıktı al."),
        ("6 - Rapor - Excel / PDF", "Öğretmen programı ve ders yükü", "Öğretmen seçip rapor al."),
        ("7 - Okul Genel Görünümü", "Tüm okul programı ve renkli dosya", "Renkli HTML/PDF için kullan."),
        ("8 - Güncelleme / Hakkında", "Sürüm, hazırlayan ve güncelleme kanalı", "Sürümü kontrol et."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)

    add_heading(doc, "3. Tanımlar ve ASC XML Aktarımı", 1)
    doc.add_paragraph("İlk gerçek kullanım genellikle ASC XML dosyasını içe aktarmakla başlar.")
    add_steps(
        doc,
        [
            "Ana menüden Tanımlar ekranını açın.",
            "ASC XML içe aktar butonuna basın.",
            "ASC XML dosyasını seçin.",
            "Önizlemede dönem, ders, öğretmen, oda, sınıf, grup, lesson ve card sayılarını kontrol edin.",
            "Onay verirseniz veriler yerel SQLite deposuna aktarılır.",
            "Aktarım sonrası sınıf, öğretmen, ders ve kaynak sekmelerinden kayıtları kontrol edin.",
        ],
    )
    add_note(doc, "Kullanıcı dosyası olan PDF veya ilgisiz dosyalar veri aktarımı sırasında kullanılmaz; ana veri kaynağı ASC XML dosyasıdır.")

    add_heading(doc, "4. Kurallar ve Öğretmen Uygunluk Kilitleri", 1)
    doc.add_paragraph(
        "Kurallar ve Kilitler ekranı öğretmenlerin ders konulamayacak kesin saatlerini veya tercih/kısıt bilgisini izlemek içindir. Eğer kısıt tanımlı değilse öğretmen satırları boş görünmek yerine durum bilgisiyle listelenir."
    )
    add_image(doc, SCREENSHOTS[3][0], f"Şekil 2 - {SCREENSHOTS[3][1]}", 6.4)

    add_heading(doc, "5. Taslak Üretme", 1)
    add_steps(
        doc,
        [
            "Ana ekranda Mevcut ASC kartlarını koru seçeneğini açık bırakın.",
            "Taslak Üret butonuna basın.",
            "Üst başlıkta atama ve yerleşmeyen talep sayılarını kontrol edin.",
            "Okul geneli görünümde bir hücrede çok ders görünmesi tek başına çakışma değildir; bu okul genelindeki aynı saat yoğunluğudur.",
            "Manuel düzenleme için sınıf veya öğretmen filtresi seçerek hücreleri tek ders seviyesine indirin.",
        ],
    )
    add_image(doc, SCREENSHOTS[0][0], f"Şekil 3 - {SCREENSHOTS[0][1]}", 6.4)

    add_heading(doc, "6. Manuel Düzenleme", 1)
    doc.add_paragraph(
        "Manuel düzenleme Taslak Çizelge ekranındadır. Sağ paneldeki Manuel düzenleme bölümü seçili dersi başka gün/saate taşımanızı veya atamayı kaldırmanızı sağlar."
    )
    add_image(doc, manual_img, "Şekil 4 - Manuel düzenleme panelinin kullanım mantığı", 6.4)
    add_steps(
        doc,
        [
            "Taslak ekranında önce Sınıf veya Öğretmen filtresi seçin.",
            "Tek ders içeren hücreye tıklayın. Sağ panelde seçili atama detayı görünür.",
            "Yeni gün ve yeni ders saati alanlarını seçin.",
            "Seçili dersi bu saate taşı butonuna basın.",
            "Uygulama sınıf, öğretmen, kaynak, blok sınırı ve kesin kilit kontrolü yapar.",
            "İşlem uygunsa değişiklik SQLite'a kaydedilir ve sonraki taslak üretiminde korunur.",
            "Seçili atamayı kaldır butonu dersi taslaktan çıkarır; Son kaldırmayı geri al ile son kaldırma geri alınabilir.",
        ],
    )
    add_note(
        doc,
        "Şu an kalıcı manuel değişiklik gün/saat taşıma ve atama kaldırma üzerinedir. Dersi başka öğretmene devretme ayrı bir geliştirme kalemidir.",
    )

    add_heading(doc, "7. Çakışma Kontrolü", 1)
    doc.add_paragraph(
        "Çakışma Kontrolü ekranı yerleşmeyen talepleri ve gerçek sınıf/öğretmen/kaynak bindirmelerini gösterir. Taslak ekranında çok ders görünen okul geneli yoğunluk bu ekran tarafından gerçek çakışma olarak sayılmayabilir."
    )
    add_image(doc, SCREENSHOTS[1][0], f"Şekil 5 - {SCREENSHOTS[1][1]}", 6.3)

    add_heading(doc, "8. Çizelgeler: Sınıf ve Kaynak/Laboratuvar Programı", 1)
    add_bullets(
        doc,
        [
            "Sınıf modu: Seçilen sınıfın haftalık ders programını gösterir.",
            "Kaynak/Lab modu: Seçilen oda, atölye veya laboratuvarın haftalık kullanım programını gösterir.",
            "CSV aktar: Excel ile açılabilecek dosya üretir.",
            "PDF/yazdır: Windows yazdırma penceresinden PDF veya yazıcı çıktısı alınır.",
        ],
    )

    add_heading(doc, "9. Öğretmen Raporu ve Ders Yükü", 1)
    doc.add_paragraph("Rapor ekranında öğretmen seçilerek haftalık program, ders yükü, sınıf öğretmenliği, kulüp, nöbet ve ek bilgiler izlenir.")
    add_image(doc, SCREENSHOTS[2][0], f"Şekil 6 - {SCREENSHOTS[2][1]}", 6.3)

    add_heading(doc, "10. Okul Genel Görünümü ve Renkli Dosya", 1)
    add_bullets(
        doc,
        [
            "Okul Genel Görünümü tüm programı filtreli şekilde incelemek içindir.",
            "Alan, program, sınıf, öğretmen ve kaynak filtreleri gerçek kurum verisinden dolar.",
            "Renkli okul dosyası çıktısı tüm okul programını HTML olarak üretir; tarayıcıdan PDF alınabilir.",
        ],
    )

    add_heading(doc, "11. ASC XML Dışa Aktarım", 1)
    add_paragraph = doc.add_paragraph
    add_paragraph(
        "Tanımlar ekranındaki ASC XML dışa aktar butonu, içe aktarılan ASC external ID bilgilerini koruyarak XML üretir. Manuel taşınan kartlarda gün/saat bilgisi güncellenir; kaldırılan kartlar çıktıya alınmaz."
    )

    add_heading(doc, "12. Güncelleme ve Hakkında", 1)
    add_bullets(
        doc,
        [
            "Menü: 8 - Güncelleme / Hakkında.",
            "Hazırlayan: Mehmet Akif SÖNMEZ.",
            "Firma: İkizsoft Bilişim Hizmetleri.",
            "Web: www.ikizsoft.com.",
            "Sürüm: Uygulama metadata'sından okunur; mevcut geliştirme sürümü 0.9.0-dev.",
            "Güncelleme kontrolü: Otomatik güncelleme sunucusu bağlanana kadar kanal bilgisi web adresi olarak gösterilir.",
        ],
    )

    add_heading(doc, "13. Hızlı Sorun Giderme", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Durum", "Yapılacak işlem"]):
        set_cell_text(table.rows[0].cells[i], h, True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    troubleshooting = [
        ("Haftalık uygunluk ızgarası boş gibi görünüyor.", "Bu ekran ders programı değil, öğretmen kilit/uygunluk ekranıdır. Kısıt yoksa öğretmenler kısıt tanımlı değil bilgisiyle listelenir."),
        ("Taslakta bir hücrede çok sayıda ders var.", "Bu okul geneli aynı saat yoğunluğudur. Sınıf veya öğretmen filtresi seçerek tek ders görünümüne geçin."),
        ("Manuel taşıma yapılmıyor.", "Önce tek ders içeren hücre seçin; hedef gün/saat seçin. Çakışma veya kesin kilit varsa uygulama taşımaz."),
        ("Çakışma kontrolü boş.", "Bu iyi durumdur: yerleşmeyen talep veya gerçek bindirme bulunmadığını gösterir."),
        ("PDF almak istiyorum.", "İlgili ekrandaki PDF/yazdır butonunu kullanın; Windows yazdırma penceresinde PDF yazıcısını seçin."),
        ("Git push yapılamıyor.", "Yerel commit var; uzak repo için origin remote tanımlanmalıdır."),
    ]
    for status, action in troubleshooting:
        cells = table.add_row().cells
        set_cell_text(cells[0], status)
        set_cell_text(cells[1], action)

    add_heading(doc, "14. Önerilen Günlük Kullanım Sırası", 1)
    add_steps(
        doc,
        [
            "ASC XML içe aktar veya mevcut veriyi kontrol et.",
            "Kurallar/Kilitler ekranında özel uygunluk durumlarını kontrol et.",
            "Taslak Üret.",
            "Sınıf veya öğretmen filtresiyle manuel gün/saat düzeltmelerini yap.",
            "Çakışma Kontrolü ekranını aç.",
            "Sınıf, kaynak/lab ve öğretmen programlarını incele.",
            "Okul Genel Görünümü üzerinden renkli okul dosyasını üret.",
            "Gerekirse CSV, PDF ve ASC XML dışa aktar.",
        ],
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Ders Dağıtım Uygulaması Kullanım Kılavuzu | İkizsoft Bilişim Hizmetleri")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(create_document())
